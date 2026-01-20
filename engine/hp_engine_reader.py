from __future__ import annotations

from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Union
from io import BytesIO, StringIO
import json
import csv
import traceback
import xml.etree.ElementTree as ET


@dataclass
class IngestedFile:
    name: str
    mime: Optional[str] = None
    ext: str = ""
    size: int = 0

    # Parsed payloads (one or more may be filled)
    text: Optional[str] = None
    tables: List[Dict[str, Any]] = field(default_factory=list)  # list of {"sheet":..., "rows":..., "cols":..., "data":...}
    json_obj: Optional[Any] = None
    xml_obj: Optional[Dict[str, Any]] = None

    # Raw fallback
    raw_bytes: Optional[bytes] = None

    # Diagnostics
    ok: bool = False
    warnings: List[str] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)


@dataclass
class IngestStore:
    files: List[IngestedFile] = field(default_factory=list)

    def summary(self) -> Dict[str, Any]:
        return {
            "total": len(self.files),
            "ok": sum(1 for f in self.files if f.ok),
            "failed": sum(1 for f in self.files if not f.ok),
            "names": [f.name for f in self.files],
        }


class HPReader:
    """
    Streamlit UploadedFile aware ingestion.
    Accepts list[UploadedFile] from st.file_uploader(accept_multiple_files=True).
    """

    def ingest(self, files: List[Any]) -> IngestStore:
        store = IngestStore()

        for uf in files or []:
            ing = IngestedFile(
                name=getattr(uf, "name", "unknown"),
                mime=getattr(uf, "type", None),
            )
            ing.ext = self._ext(ing.name)
            try:
                b = uf.getvalue() if hasattr(uf, "getvalue") else bytes(uf)
                ing.raw_bytes = b
                ing.size = len(b)
                self._parse_into(ing, b)
            except Exception as e:
                ing.ok = False
                ing.errors.append(f"Ingest crash: {e}")
                ing.errors.append(traceback.format_exc())

            store.files.append(ing)

        return store

    # -------------------------
    # Internals
    # -------------------------

    def _ext(self, filename: str) -> str:
        fn = (filename or "").lower()
        if "." not in fn:
            return ""
        return fn.rsplit(".", 1)[-1]

    def _parse_into(self, ing: IngestedFile, b: bytes) -> None:
        ext = ing.ext

        # --- TEXT-LIKE ---
        if ext in {"txt", "md", "log"}:
            ing.text = self._decode_text(b, ing)
            ing.ok = ing.text is not None
            return

        if ext in {"csv"}:
            text = self._decode_text(b, ing)
            if text is None:
                ing.ok = False
                return
            ing.text = text
            ing.tables.append(self._parse_csv_table(text, ing))
            ing.ok = True
            return

        if ext in {"json"}:
            text = self._decode_text(b, ing)
            if text is None:
                ing.ok = False
                return
            ing.text = text
            try:
                ing.json_obj = json.loads(text)
                ing.ok = True
            except Exception as e:
                ing.ok = False
                ing.errors.append(f"JSON parse error: {e}")
            return

        if ext in {"xml"}:
            text = self._decode_text(b, ing)
            if text is None:
                ing.ok = False
                return
            ing.text = text
            try:
                ing.xml_obj = self._parse_xml(text)
                ing.ok = True
            except Exception as e:
                ing.ok = False
                ing.errors.append(f"XML parse error: {e}")
            return

        # --- EXCEL ---
        if ext in {"xlsx", "xls"}:
            try:
                import pandas as pd  # type: ignore
            except Exception:
                ing.warnings.append("pandas not available; keeping raw bytes only.")
                ing.ok = True  # store raw
                return

            # xls needs xlrd; xlsx needs openpyxl. We'll attempt and warn on missing.
            bio = BytesIO(b)
            try:
                if ext == "xlsx":
                    # openpyxl required under the hood
                    xls = pd.ExcelFile(bio, engine="openpyxl")
                else:
                    # xlrd required for old xls
                    xls = pd.ExcelFile(bio)
            except Exception as e:
                ing.warnings.append(f"Excel engine missing or file unreadable: {e}. Keeping raw bytes.")
                ing.ok = True
                return

            for sheet in xls.sheet_names:
                try:
                    df = xls.parse(sheet)
                    ing.tables.append({
                        "sheet": sheet,
                        "rows": int(df.shape[0]),
                        "cols": int(df.shape[1]),
                        "data": df.fillna("").to_dict(orient="records"),
                    })
                except Exception as e:
                    ing.warnings.append(f"Sheet '{sheet}' read failed: {e}")

            ing.ok = True
            return

        # --- DOCX ---
        if ext in {"docx"}:
            try:
                import docx  # python-docx
            except Exception:
                ing.warnings.append("python-docx not available; keeping raw bytes only.")
                ing.ok = True
                return

            try:
                d = docx.Document(BytesIO(b))
                paras = [p.text for p in d.paragraphs if p.text and p.text.strip()]
                ing.text = "\n".join(paras) if paras else ""
                ing.ok = True
            except Exception as e:
                ing.ok = False
                ing.errors.append(f"DOCX read error: {e}")
            return

        # --- PDF ---
        if ext in {"pdf"}:
            # Use pypdf if installed; otherwise store raw with warning.
            try:
                from pypdf import PdfReader  # type: ignore
            except Exception:
                ing.warnings.append("pypdf not available; keeping raw bytes only.")
                ing.ok = True
                return

            try:
                reader = PdfReader(BytesIO(b))
                pages_text = []
                for i, page in enumerate(reader.pages):
                    try:
                        t = page.extract_text() or ""
                        if t.strip():
                            pages_text.append(t)
                    except Exception as e:
                        ing.warnings.append(f"PDF page {i} extract failed: {e}")
                ing.text = "\n\n".join(pages_text)
                ing.ok = True
            except Exception as e:
                ing.ok = False
                ing.errors.append(f"PDF read error: {e}")
            return

        # --- FALLBACK ---
        ing.warnings.append(f"Unsupported extension '{ext}'. Keeping raw bytes only.")
        ing.ok = True

    def _decode_text(self, b: bytes, ing: IngestedFile) -> Optional[str]:
        # Try UTF-8, then Windows-1254 (TR), then latin-1
        for enc in ("utf-8", "utf-8-sig", "cp1254", "latin-1"):
            try:
                return b.decode(enc)
            except Exception:
                continue
        ing.errors.append("Text decode failed for utf-8/cp1254/latin-1")
        return None

    def _parse_csv_table(self, text: str, ing: IngestedFile) -> Dict[str, Any]:
        # Very tolerant CSV parser
        sio = StringIO(text)
        try:
            reader = csv.reader(sio)
            rows = list(reader)
            if not rows:
                return {"sheet": "csv", "rows": 0, "cols": 0, "data": []}
            header = rows[0]
            data_rows = rows[1:]
            # Map to dict records
            records = []
            for r in data_rows:
                rec = {}
                for i, h in enumerate(header):
                    rec[h] = r[i] if i < len(r) else ""
                records.append(rec)
            return {"sheet": "csv", "rows": len(records), "cols": len(header), "data": records}
        except Exception as e:
            ing.warnings.append(f"CSV parse warning: {e}")
            return {"sheet": "csv", "rows": 0, "cols": 0, "data": []}

    def _parse_xml(self, text: str) -> Dict[str, Any]:
        root = ET.fromstring(text)

        def node_to_dict(node: ET.Element) -> Dict[str, Any]:
            d: Dict[str, Any] = {"tag": node.tag}
            if node.attrib:
                d["attributes"] = dict(node.attrib)
            if node.text and node.text.strip():
                d["text"] = node.text.strip()
            children = list(node)
            if children:
                d["children"] = [node_to_dict(c) for c in children]
            return d

        return node_to_dict(root)